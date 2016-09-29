# coding=utf8
from docx import Document
from config import PATH

def createDocx(list, name, num,  grade='2014'):
    document = Document()
    document.add_heading(u'成绩详细信息', 0)
    document.add_heading(u'姓名:'+name + u"\n学号:"+num, 1)
    document.add_paragraph(u'(以下是所有参与平均分计算的课程)')
    document.add_paragraph(u'\n必修部分')

    table = document.add_table(rows=1, cols=4,style = 'Table Grid')
    hd_cells = table.rows[0].cells
    hd_cells[0].text = u'序号'
    hd_cells[1].text = u'课程名称'
    hd_cells[2].text = u'学分'
    hd_cells[3].text = u'成绩'

    cnt = 1
    for i in list:
        if u"必修" in i["type"] and float(i["grade"]) != 0 \
                and i["year"] == "2015" and (i["putong"] == u"普通" or i["putong"] == u"重修") and i["used"] == True:
            if grade == "2015" and i["apart"] == u"体育部":
                continue
            else:
                row_cells = table.add_row().cells
                row_cells[0].text = str(cnt)
                row_cells[1].text = i["name"]
                row_cells[2].text = i["point"]
                row_cells[3].text = str(i["grade"])
                cnt += 1

    document.add_paragraph(u'\n选修部分')
    table2 = document.add_table(rows=1, cols=4,style = 'Table Grid')
    hd_cells = table2.rows[0].cells
    hd_cells[0].text = u'序号'
    hd_cells[1].text = u'课程名称'
    hd_cells[2].text = u'学分'
    hd_cells[3].text = u'成绩'


    cnt = 1
    for i in list:
        if (u"选修" in i['type'] or u"辅修" in i["putong"]) and i["year"] == "2015" and i["used"] == True:
            row_cells = table2.add_row().cells
            row_cells[0].text = str(cnt)
            row_cells[1].text = i["name"]
            row_cells[2].text = i["point"]
            row_cells[3].text = str(i["grade"])
            cnt += 1

    document.save(PATH + u"app/temp/file_doc/成绩详情_{}.docx".format(str(num)))


